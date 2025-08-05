import streamlit as st
from modules.generate_text import (
    generate_diagnosis_summary,
    suggest_goals,
    generate_emergency_plan,
    write_background,
    rewrite_interventions,
    refine_text
)
from modules.file_parser import parse_document
from modules.location_utils import get_coordinates_from_address, find_nearest_park
from docx import Document
import io

st.set_page_config(page_title="Care Plan Generator", layout="wide")
st.title("🧾 Nurse Next Door – Care Plan Generator")

if "final_plan" not in st.session_state:
    st.session_state.final_plan = ""
if "editable_plan" not in st.session_state:
    st.session_state.editable_plan = ""

responses = {}
intake_type = st.radio("Select Intake Source", ["Organic", "Support Coordinator", "Carer's Gateway"])

# Common Fields
def editable_fields():
    responses["Client Name"] = st.text_input("Client Full Name", value=responses.get("Client Name", ""))
    responses["Sex"] = st.selectbox("Sex", ["", "Male", "Female", "Other"], index=["", "Male", "Female", "Other"].index(responses.get("Sex", "")))
    responses["Date of Birth"] = st.date_input("Date of Birth")
    responses["Emergency Contact Name"] = st.text_input("Emergency Contact Name", value=responses.get("Emergency Contact Name", ""))
    responses["Emergency Contact Phone"] = st.text_input("Emergency Contact Phone", value=responses.get("Emergency Contact Phone", ""))
    responses["Emergency Contact Email"] = st.text_input("Emergency Contact Email", value=responses.get("Emergency Contact Email", ""))
    responses["Client Address"] = st.text_input("Client Address", value=responses.get("Client Address", ""))
    responses["Preferred Language"] = st.selectbox("Preferred Language", ["English", "Hindi", "Punjabi", "Mandarin", "Other"])
    return responses

# ================= ORGANIC =================
if intake_type == "Organic":
    st.subheader("📋 Client Information")
    editable_fields()

    st.subheader("🧠 Medical & Care Info")
    diagnosis_input = st.text_area("Main Diagnosis (brief)", height=80)
    services_required = st.text_area("Services Required", height=80)
    allergies = st.text_input("Allergies")
    alerts = st.text_input("Alerts / Risks")
    medications = st.text_area("Current Medications")
    mobility = st.text_area("Mobility Needs", height=100)
    interventions = st.text_area("Raw Intervention Notes (will be GPT-enhanced)", height=120)

    st.subheader("🏡 Background")
    lifestyle_info = st.text_area("Describe lifestyle, culture, support systems, and interests")

    if st.button("✨ Generate Enhanced Care Plan Elements"):
        with st.spinner("Generating enhanced care plan components..."):
            client_name = responses["Client Name"]
            evac_point = "a local park"
            lat, lon = get_coordinates_from_address(responses["Client Address"])
            if lat and lon:
                evac_point = find_nearest_park(lat, lon)

            responses["Diagnosis Summary"] = generate_diagnosis_summary(diagnosis_input, client_name)
            responses["Goals"] = suggest_goals(diagnosis_input, services_required, client_name)
            responses["Emergency Plan"] = generate_emergency_plan(
                client_name, evac_point, responses["Preferred Language"],
                responses["Emergency Contact Name"], responses["Emergency Contact Phone"]
            )
            responses["Background"] = write_background(client_name, lifestyle_info)
            responses["Interventions"] = rewrite_interventions(
                interventions, alerts, allergies, medications, diagnosis_input, client_name
            )

            st.session_state.final_plan = f"""
**Client:** {client_name}

**Diagnosis Summary:**  
{responses["Diagnosis Summary"]}

**Support Goals:**  
{responses["Goals"]}

**Emergency Preparedness Plan:**  
{responses["Emergency Plan"]}

**Background:**  
{responses["Background"]}

**Care Interventions:**  
{responses["Interventions"]}
"""
            st.session_state.editable_plan = st.session_state.final_plan

        st.success("Enhanced care plan generated.")

# ================= FILE-BASED =================
elif intake_type in ["Support Coordinator", "Carer's Gateway"]:
    st.subheader("📎 Upload Referral Files")
    uploaded_files = st.file_uploader("Upload PDF or DOCX files", type=["pdf", "docx"], accept_multiple_files=True)

    if uploaded_files:
        full_text = ""
        for uploaded_file in uploaded_files:
            parsed = parse_document(uploaded_file)
            full_text += parsed.get("Raw Text", "") + "\n"

        responses["Raw Text"] = full_text
        st.write("📄 Parsed Text:")
        st.text_area("Editable Text", value=full_text, height=300)

    st.subheader("📋 Complete Missing Client Info")
    editable_fields()

    diagnosis_input = st.text_input("What is the main diagnosis you extracted from this file?")
    services_required = st.text_input("What services are likely required?")
    interventions = st.text_area("Any additional notes that could help in writing interventions")
    lifestyle_info = st.text_area("Lifestyle/Background Notes")

    if st.button("✨ Generate Care Plan from File"):
        with st.spinner("Generating care plan from referral..."):
            client_name = responses["Client Name"]
            evac_point = "a local park"
            lat, lon = get_coordinates_from_address(responses["Client Address"])
            if lat and lon:
                evac_point = find_nearest_park(lat, lon)

            # Fall back in case these are not collected earlier
            allergies = responses.get("Allergies", "")
            alerts = responses.get("Alerts / Risks", "")
            medications = responses.get("Current Medications", "")

            responses["Diagnosis Summary"] = generate_diagnosis_summary(diagnosis_input, client_name)
            responses["Goals"] = suggest_goals(diagnosis_input, services_required, client_name)
            responses["Emergency Plan"] = generate_emergency_plan(
                client_name, evac_point, responses["Preferred Language"],
                responses["Emergency Contact Name"], responses["Emergency Contact Phone"]
            )
            responses["Background"] = write_background(client_name, lifestyle_info)
            responses["Interventions"] = rewrite_interventions(
                interventions or responses.get("Raw Text", ""), alerts, allergies, medications, diagnosis_input, client_name
            )

            st.session_state.final_plan = f"""
**Client:** {client_name}

**Diagnosis Summary:**  
{responses["Diagnosis Summary"]}

**Support Goals:**  
{responses["Goals"]}

**Emergency Preparedness Plan:**  
{responses["Emergency Plan"]}

**Background:**  
{responses["Background"]}

**Care Interventions:**  
{responses["Interventions"]}
"""
            st.session_state.editable_plan = st.session_state.final_plan

        st.success("Care plan generated from referral.")

# ================= OUTPUT & FEEDBACK =================
if st.session_state.final_plan:
    st.subheader("📄 Final Care Plan")

    st.session_state.editable_plan = st.text_area("📋 View/Edit Care Plan", value=st.session_state.editable_plan, height=700)
    feedback = st.text_area("💬 Provide overall feedback (optional)", key="global_feedback")

    if st.button("🔁 Regenerate Based on Feedback"):
        with st.spinner("Applying feedback to care plan..."):
            refined = refine_text(st.session_state.editable_plan, feedback)
            st.session_state.final_plan = refined
            st.session_state.editable_plan = refined
            st.rerun()

    # =========== DOWNLOAD =============
    doc = Document()
    doc.add_heading("Nurse Next Door – Care Plan", 0)
    for line in st.session_state.final_plan.strip().split("\n"):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📥 Download Care Plan as Word Document",
        data=buffer,
        file_name=f"{responses.get('Client Name','care_plan')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
